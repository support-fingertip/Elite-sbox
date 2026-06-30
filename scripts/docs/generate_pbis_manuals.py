#!/usr/bin/env python3
"""
Generates three Microsoft Word user manuals for the Primary PBIS module
(plus the Secondary PBIS KPI), one per audience:

  docs/Primary_PBIS_Admin_Manual.docx
  docs/Primary_PBIS_Sales_User_Manual.docx
  docs/Primary_PBIS_Hierarchy_User_Manual.docx

Run:  python3 scripts/docs/generate_pbis_manuals.py
"""

import os
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "docs")
COMPANY = "Elite Foods"
MODULE = "Primary PBIS (Performance Based Incentive Scheme)"
VERSION = "1.0"
TODAY = date.today().strftime("%d %B %Y")

BRAND = RGBColor(0x0B, 0x5C, 0xAB)
BRAND_DARK = RGBColor(0x01, 0x44, 0x86)
GREY = RGBColor(0x70, 0x6E, 0x6B)


# ---------- low-level helpers ----------
def _shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _base_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for i, sz in [(1, 18), (2, 14), (3, 12)]:
        st = doc.styles["Heading %d" % i]
        st.font.name = "Calibri"
        st.font.size = Pt(sz)
        st.font.color.rgb = BRAND if i == 1 else BRAND_DARK


def title_page(doc, audience):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(COMPANY)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = BRAND

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(MODULE)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = BRAND_DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("User Manual")
    r.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(audience)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = BRAND

    for _ in range(8):
        doc.add_paragraph()
    for line in ["Version %s" % VERSION, "Last updated: %s" % TODAY,
                 "Applies to: Salesforce (Lightning Experience) — Desktop & Mobile"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p.add_run(line)
        rr.font.color.rgb = GREY
    doc.add_page_break()


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def h3(doc, text):
    doc.add_heading(text, level=3)


def para(doc, text):
    return doc.add_paragraph(text)


def steps(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def callout(doc, label, text, fill="EAF3FF"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    _shade(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(label + ": ")
    r.bold = True
    r.font.color.rgb = BRAND_DARK
    p.add_run(text)
    doc.add_paragraph()


def screenshot(doc, caption):
    p = doc.add_paragraph()
    r = p.add_run("[ Screenshot placeholder — %s ]" % caption)
    r.italic = True
    r.font.color.rgb = GREY


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].paragraphs[0].add_run(htext).bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


# ---------- shared content ----------
def intro_what_is(doc):
    h1(doc, "1. What is Primary PBIS?")
    para(doc,
         "Primary PBIS is Elite Foods' configurable engine for calculating Performance Based "
         "Incentives on primary sales. For every salesperson and period it measures actual "
         "achievement against targets across a set of Sales & Distribution (S&D) parameters, "
         "matches each achievement to an incentive slab, and computes the incentive earned.")
    h3(doc, "The building blocks")
    table(doc,
          ["Component", "What it is"],
          [["S & D Parameter", "A reusable measure (e.g. Total Volume, Focus Pack, New Customers). "
            "Defines what is measured and how (the operator)."],
           ["Focused Pack", "A named set of SKUs used by Focus-Pack parameters, scoped by Sales Channel."],
           ["PBIS Policy", "A channel-wise policy holding, per Profile, the incentive slabs "
            "(achievement bands and their Monthly/Quarterly incentive %)."],
           ["PBIS Target", "One per user per period. Holds up to 10 S&D parameter slots with each "
            "slot's Target and Weightage, and (after calculation) the Achievement, %, Slab and Incentive."],
           ["KPI Dashboard", "A tab with two sub-tabs — Primary KPI and Secondary KPI — where users "
            "and managers view performance and incentives."]])
    callout(doc, "How incentive is calculated",
            "For each filled slot: Incentive = Gross Salary x Slab % x Weightage. "
            "The Total Incentive is paid only when the target is Eligible (every mandatory parameter "
            "reaches the lowest slab floor); otherwise the total is 0.")


# ---------- ADMIN ----------
def build_admin():
    doc = Document()
    _base_styles(doc)
    title_page(doc, "Administrator Guide")

    intro_what_is(doc)

    h1(doc, "2. Before you begin")
    bullets(doc, [
        "You need an admin profile (or a permission set granting access to S & D Parameters, "
        "Primary PBIS Policies, PBIS Targets, Focused Pack and the KPI Dashboard).",
        "Each salesperson is a Salesforce User with a Role (the role hierarchy powers the team views) "
        "and a Monthly Gross Salary on their User record.",
        "Setup order: (1) Focused Packs, (2) S & D Parameters, (3) PBIS Policy + slabs, "
        "(4) PBIS Targets, (5) run the calculation.",
    ])

    h1(doc, "3. Create Focused Packs")
    para(doc, "Focused Packs group the SKUs that a Focus-Pack parameter measures, per Sales Channel.")
    steps(doc, [
        "Open the Focused Pack tab and click New.",
        "Enter the Focused Pack Name and choose the Sales Channel.",
        "Use the filters (Category / Group / Sub Group / Search) and the All / Selected toggle to "
        "find and tick the SKUs that belong to this pack.",
        "Click Save. The pack now appears for parameters in that same Sales Channel.",
    ])
    callout(doc, "Tip", "The Focused Pack picker on an S&D Parameter only lists packs whose Sales "
            "Channel matches the parameter's channel — set the channel first.")
    screenshot(doc, "Focused Pack builder with SKU list and All/Selected toggle")

    h1(doc, "4. Create S & D Parameters")
    para(doc, "An S&D Parameter (tab: S & D Parameters) defines one measure. Click New to open the "
              "full-screen builder, choose an Operator, then fill the operator-specific fields.")
    h3(doc, "Operators")
    table(doc,
          ["Operator", "Use it for", "Key fields"],
          [["SUM", "Sum a numeric field (e.g. volume MT, revenue)", "Object, Field, User Field, Date Field, Filters"],
           ["COUNT", "Count records (e.g. orders)", "Object, User Field, Date Field, Filters"],
           ["COUNT_DISTINCT", "Unique count of one field (e.g. unique Accounts on Order)", "Object, Field (to count), User Field, Date Field"],
           ["MULTI_OBJECT_DISTINCT", "Unique Customers across several objects (e.g. Visit + Visit Form + Order)",
            "One or more Sources, each: Object, Customer Field, User Field, Date Field, Filters"],
           ["FOCUS_PACK_VOLUME / FOCUS_PACK_REVENUE", "Volume/revenue restricted to a Focused Pack's SKUs",
            "Object, Field, User/Date Field, Sales Channel, Focused Pack"]])
    h3(doc, "Steps")
    steps(doc, [
        "S & D Parameters tab > New.",
        "Enter a Name and pick the Operator.",
        "Choose the source Object; pick the User Field (identifies the salesperson) and the Date Field "
        "(used for the period).",
        "For SUM / COUNT_DISTINCT / Focus Pack: pick the Field to aggregate/count.",
        "For Unique Customers (multi-object): click Add Source and define each object + its customer "
        "field (+ user/date field + filters). Add a source per customer field (e.g. Visit Form Primary "
        "and Secondary Customer are two sources).",
        "Optionally add Filters and Filter Logic; tick Mandatory if this is a value/volume gate.",
        "Click Save. The query the engine will run is stored automatically.",
    ])
    callout(doc, "Important", "Mark at least one parameter per target as Mandatory — eligibility "
            "depends on the mandatory parameter(s) reaching the lowest slab floor.")
    callout(doc, "After deployment", "If parameters were just deployed, an admin runs "
            "scripts/apex/regen_sd_parameter_soql.apex once so stored queries are regenerated.")
    screenshot(doc, "S&D Parameter builder showing operator and source mapping")

    h1(doc, "5. Build a PBIS Policy and its Slabs")
    steps(doc, [
        "Primary PBIS Policies tab > New to open the policy builder.",
        "Enter the Policy Name, Basis (Value/Volume), and select the Sales Channel(s).",
        "Add a slab row per Profile band: Profile, From %, To %, Monthly % and Quarterly %.",
        "Use Add Slab for more bands; leave the top band's To % blank for an open-ended top slab.",
        "Click Save.",
    ])
    callout(doc, "One active policy per channel",
            "An active policy may not share a Sales Channel with another active policy. If you try, "
            "you'll see an error naming the existing policy (e.g. \"KA Policy\" (KA)). Deactivate the "
            "old one or change the channel.")
    screenshot(doc, "PBIS Policy builder with per-Profile slab rows")

    h1(doc, "6. Upload PBIS Targets")
    para(doc, "PBIS Targets are loaded per user per period using Salesforce Inspector (Data Import). "
              "A template is provided at scripts/import/PBIS_Target_Import_Template.xlsx.")
    h3(doc, "Template & mapping")
    bullets(doc, [
        "Row 1 = field API names; Row 2 = descriptions; Row 3 = a sample. Remove rows 2-3 before importing.",
        "Operation: Upsert, External Id field: External_Id__c (so re-imports update the same row).",
        "User by Username: column User__r.Username. Policy & Profile derive from the User; Gross Salary "
        "auto-fills from the User on save.",
        "Each slot's parameter: column S_D_Parameter_<n>__r.External_Id__c (the parameter's External Id).",
        "Per slot also provide S_D_Parameter_<n>_Weightage__c and S_D_Parameter_<n>_Target__c.",
    ])
    h3(doc, "Integrity rules enforced on save")
    table(doc,
          ["Rule", "Message you may see", "Fix"],
          [["Slot all-or-none", "S & D Parameter N: enter the Parameter, its Target and its Weightage "
            "together, or leave all three blank.", "Fill all three for the slot, or clear all three."],
           ["Weightages total 100", "S & D Parameter weightages must total 100%.", "Adjust weightages to sum to 100."],
           ["Mandatory required", "Include the mandatory value/volume S&D parameter in at least one slot.",
            "Add a parameter marked Mandatory."],
           ["End Date in future", "For new PBIS Targets, End Date must be later than today.", "Use a current/future period."],
           ["No overlapping period", "A PBIS Target already exists for this User and Incentive Period with an "
            "overlapping period ...", "Use non-overlapping dates, or a different Incentive Period."],
           ["Unique parameters", "S & D Parameter 1 and S & D Parameter 2 are matching. Please add unique "
            "Parameters.", "Use a different S&D Parameter in each slot."]])
    callout(doc, "Bulk uploads", "All rules are bulk-safe for ~1000 users. With partial-success enabled, "
            "only the offending rows fail; the rest load. Targets are owned by their User automatically.")
    screenshot(doc, "Salesforce Inspector Data Import — Upsert on External_Id__c")

    h1(doc, "7. Run & verify the calculation")
    para(doc, "Actuals and incentive are computed from each parameter's stored query and the policy slabs. "
              "There are three ways to run it:")
    table(doc,
          ["Method", "Where", "Scope"],
          [["Recalculate Incentive", "Quick action on a PBIS Target record", "That one record (immediate)"],
           ["Refresh Actuals", "List button on the PBIS Targets list view", "All active in-period targets (background)"],
           ["Daily scheduler", "Automatic at 9:00 AM", "All active targets whose period includes today"]])
    para(doc, "After running, each slot shows Achievement, Achievement %, the qualified Slab band, "
              "Incentive % and Incentive Amount; the header shows Total Incentive, Max Possible and Eligible.")
    callout(doc, "Slab snapshot", "Each slot also stores S_D_Parameter_N_Slab__c (the band, e.g. "
            "100 - 109.99) and S_D_Parameter_N_Incentive_Percentage__c, so the record keeps a trace even "
            "if the slab is later changed or deleted.")

    h1(doc, "8. KPI Dashboard & access")
    bullets(doc, [
        "The KPI Dashboard tab contains two sub-tabs: Primary KPI and Secondary KPI.",
        "Grant the KPI Dashboard tab and Apex access via the KPI_Dashboard_Access permission set "
        "(and add the tab to the relevant Lightning app).",
        "Ensure users have field-level security (via a profile or permission set) on the PBIS Target "
        "fields so the dashboard and record page show values.",
        "The record page Primary PBIS Target Record Page shows the per-slot Slab band and Incentive % "
        "(only on saved records).",
    ])

    h1(doc, "9. Deployment notes")
    bullets(doc, [
        "Deploy via the provided manifests/workflows (e.g. pbis_kpi_slab.xml, deploy-pbisactuals).",
        "The KPI container references the existing Secondary KPI dashboard component — it must exist "
        "in the target org.",
        "After deploy: run scripts/apex/regen_sd_parameter_soql.apex, then re-run a recalc so existing "
        "targets populate the new fields.",
        "Schedule the daily job once: scripts/apex/schedule_pbis_actuals.apex.",
    ])

    h1(doc, "10. Troubleshooting")
    table(doc,
          ["Symptom", "Likely cause / fix"],
          [["Incentive is 0 although achievement is high", "Target is Not Eligible — a mandatory parameter "
            "is below the lowest slab floor. Check the mandatory slot's %."],
           ["KPI dashboard columns blank for a user", "Missing field-level security on PBIS Target fields."],
           ["Achievement shows 0 after import", "Recalc not run yet, or parameter query needs regeneration "
            "(run the regen script) — then Refresh Actuals."],
           ["Cannot save a policy (channel error)", "Another active policy already uses that channel."],
           ["Row rejected on import", "See the integrity-rules table in section 6."]])

    out = os.path.join(OUT_DIR, "Primary_PBIS_Admin_Manual.docx")
    doc.save(out)
    return out


# ---------- SALES USERS ----------
def build_sales():
    doc = Document()
    _base_styles(doc)
    title_page(doc, "Sales User Guide")

    intro_what_is(doc)

    h1(doc, "2. Opening your KPI Dashboard")
    steps(doc, [
        "Click the App Launcher (the grid icon) or your app's navigation bar.",
        "Open the KPI Dashboard tab.",
        "Make sure the Primary KPI sub-tab is selected (the second sub-tab, Secondary KPI, is covered "
        "in section 6).",
    ])
    screenshot(doc, "KPI Dashboard tab with Primary KPI / Secondary KPI switch")

    h1(doc, "3. Choosing the period")
    para(doc, "At the top of Primary KPI, set the period you want to view:")
    bullets(doc, [
        "Year and Month — the month you want to see.",
        "Period — Monthly or Quarterly (which kind of target to show).",
        "Use Refresh to reload if needed.",
    ])
    callout(doc, "Note", "The dashboard reads the latest calculated values. Numbers update after the "
            "nightly calculation or when an admin recalculates.")

    h1(doc, "4. Reading your headline cards")
    para(doc, "The coloured cards summarise your period at a glance:")
    table(doc,
          ["Card", "Meaning"],
          [["Total Incentive", "The incentive you have earned this period (0 if you are Not Eligible)."],
           ["Achievement", "Your overall achievement %, weighted across your parameters."],
           ["Eligibility", "Eligible or Not eligible — whether the mandatory target is met."]])
    screenshot(doc, "Primary KPI hero cards: Total Incentive, Achievement, Eligibility")

    h1(doc, "5. The S & D Parameter breakdown")
    para(doc, "Below the cards, the breakdown shows each parameter on your target:")
    table(doc,
          ["Column", "Meaning"],
          [["Parameter", "The S&D parameter name."],
           ["Weight %", "Its weightage toward your incentive."],
           ["Target / Achievement", "Your target and what you actually achieved."],
           ["Ach %", "Achievement as a percentage of target."],
           ["Slab", "The qualified slab band, e.g. 100 - 109.99."],
           ["Incentive %", "The slab's incentive percentage applied."],
           ["Incentive", "The incentive earned from this parameter."]])
    callout(doc, "On mobile", "Each parameter appears as a compact card with a progress bar — just "
            "scroll; there is no wide table to swipe.")
    screenshot(doc, "Parameter breakdown table (desktop) / cards (mobile)")

    h1(doc, "6. Understanding Eligibility & incentive")
    bullets(doc, [
        "Incentive per parameter = Gross Salary x Slab % x Weightage.",
        "Your Total Incentive is paid only when you are Eligible — i.e. the mandatory (value/volume) "
        "parameter reaches the lowest slab band.",
        "If you are Not eligible, parameter incentives may still be shown for reference but the Total is 0.",
        "Aim to push the mandatory parameter above the entry slab to unlock the total.",
    ])

    h1(doc, "7. Secondary KPI")
    para(doc, "The Secondary KPI sub-tab shows your secondary-sales KPIs (targets, achievement and "
              "secondary incentives). Switch to it from the toggle at the top of the KPI Dashboard; "
              "pick the Year/Month there to view a period.")
    screenshot(doc, "Secondary KPI dashboard")

    h1(doc, "8. FAQ")
    table(doc,
          ["Question", "Answer"],
          [["My numbers look old", "They refresh after the nightly run; ask your admin to recalculate "
            "if you need them sooner."],
           ["A parameter shows 0 achievement", "No qualifying activity was recorded for that parameter "
            "in the period, or the period/filters exclude it."],
           ["Total Incentive is 0", "You are Not eligible — check the Eligibility card and the mandatory "
            "parameter's Ach %."],
           ["Can I see a teammate's numbers?", "Only managers see team members (see the Hierarchy User guide)."]])

    out = os.path.join(OUT_DIR, "Primary_PBIS_Sales_User_Manual.docx")
    doc.save(out)
    return out


# ---------- HIERARCHY USERS ----------
def build_hierarchy():
    doc = Document()
    _base_styles(doc)
    title_page(doc, "Hierarchy / Manager Guide")

    intro_what_is(doc)

    h1(doc, "2. Your views: My, My Team, User Search")
    para(doc, "On Primary KPI, managers get a View selector with three options:")
    table(doc,
          ["View", "Shows"],
          [["My", "Your own Primary PBIS (same as a sales user)."],
           ["My Team", "A summary of your team plus member and breakdown tables."],
           ["User Search", "Search and open any one team member's Primary PBIS."]])
    callout(doc, "Who is my team?", "Your team is everyone in the roles below yours in the Salesforce "
            "role hierarchy. If you have no subordinate roles, you only see the My view.")
    screenshot(doc, "Primary KPI with Year / Month / Period / View controls")

    h1(doc, "3. The My Team view")
    h3(doc, "Team summary cards")
    bullets(doc, [
        "Total Team Incentive — the sum of your team members' incentives for the period.",
        "Achievement — the team's average achievement %.",
    ])
    h3(doc, "Team members table")
    para(doc, "One row per member with Name, Role, Achievement %, Incentive and Eligible. "
              "Use the row action View this user (or tap a card on mobile) to drill into that person.")
    h3(doc, "Breakdowns")
    table(doc,
          ["Panel", "What it shows"],
          [["Sales Channel wise", "Per Sales Channel: average Achievement % and total Incentive earned "
            "across the team."],
           ["S & D Parameter wise", "Per S&D parameter (across all members' slots): average Achievement % "
            "and total Incentive earned."]])
    screenshot(doc, "My Team: summary cards, members table, Sales-Channel-wise & S&D-Parameter-wise panels")

    h1(doc, "4. Viewing one team member (drill-down / search)")
    steps(doc, [
        "From My Team, click View this user on a member row (or tap the member card on mobile).",
        "The dashboard switches to that person's Primary PBIS — headline cards and the full parameter "
        "breakdown (with Slab band and Incentive %).",
        "Click Back to team to return to the team view.",
        "Alternatively choose View = User Search and type a name to open any team member directly.",
    ])
    callout(doc, "Heading", "In User Search / drill-down the heading shows the selected person's name. "
            "If they have no target for the period you'll see a clear 'no Primary PBIS targets' message.")
    screenshot(doc, "Drill-down into a team member's Primary PBIS")

    h1(doc, "5. Acting on team performance")
    bullets(doc, [
        "Use Sales Channel wise to spot channels lagging on achievement.",
        "Use S & D Parameter wise to see which parameters drive (or miss) incentive across the team.",
        "Drill into Not eligible members to see which mandatory parameter is short of the entry slab.",
        "Change Year/Month/Period to compare across periods.",
    ])

    h1(doc, "6. Secondary KPI (team)")
    para(doc, "The Secondary KPI sub-tab provides the secondary-sales team dashboard: your DSM/SSA team, "
              "achievement by criterion and by channel, and top / needs-attention performers, with the "
              "same drill-down idea. Switch to it from the toggle at the top of the KPI Dashboard.")
    screenshot(doc, "Secondary KPI team dashboard")

    h1(doc, "7. FAQ")
    table(doc,
          ["Question", "Answer"],
          [["I don't see My Team", "You have no subordinate roles, or your role isn't above others in the "
            "hierarchy. Ask your admin to check the role hierarchy."],
           ["A member is missing from the team", "They may have no target for the selected period, or are "
            "inactive, or sit outside your role branch."],
           ["Team totals look low", "Members may be Not eligible (incentive 0) — drill in to confirm."],
           ["Mobile use", "All team tables become scrollable cards; the views work on the Salesforce mobile app."]])

    out = os.path.join(OUT_DIR, "Primary_PBIS_Hierarchy_User_Manual.docx")
    doc.save(out)
    return out


if __name__ == "__main__":
    os.makedirs(OUT_DIR, exist_ok=True)
    for fn in (build_admin, build_sales, build_hierarchy):
        print("Generated:", os.path.normpath(fn()))
