#!/usr/bin/env python3
"""Generate docs/Schemes-User-Manual.docx from the manual content.

Self-contained (does not parse the Markdown) so the Word formatting (headings,
styled tables, shaded header rows) is fully controlled.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x02, 0x32, 0x3E)
GREEN = RGBColor(0x15, 0x80, 0x3D)
GREY = RGBColor(0x6B, 0x72, 0x80)

doc = Document()

# Base font
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

for s in ("Heading 1", "Heading 2", "Heading 3"):
    st = doc.styles[s]
    st.font.color.rgb = NAVY
    st.font.name = "Calibri"


def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        shade(hdr[i], "02323E")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def para(text="", *, italic=False, bold=False, color=None, size=11, space_after=6):
    p = doc.add_paragraph()
    if text:
        r = p.add_run(text)
        r.italic = italic
        r.bold = bold
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullets(items, style="List Bullet"):
    for it in items:
        p = doc.add_paragraph(style=style)
        # support a leading "**bold** rest" pattern
        if it.startswith("**") and "**" in it[2:]:
            bold_part, rest = it[2:].split("**", 1)
            rb = p.add_run(bold_part)
            rb.bold = True
            p.add_run(rest)
        else:
            p.add_run(it)


def numbered(items):
    bullets(items, style="List Number")


# ----------------------------------------------------------------------------- title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
tr = title.add_run("Schemes — User Manual")
tr.bold = True
tr.font.size = Pt(24)
tr.font.color.rgb = NAVY

para(
    "This manual explains how to set up and use Schemes in the system, end to end, as three "
    "process-step parts: Part 1 — Scheme Product Creation, Part 2 — Scheme Creation and "
    "Applicability (including the scheme types with a worked example for each), and Part 3 — "
    "Creation of Order. A short concepts overview and a troubleshooting checklist are included.",
    color=GREY,
)

# ----------------------------------------------------------------------------- 1
doc.add_heading("1. Overview & Key Concepts", level=1)
para(
    "A Scheme is a promotional offer (free quantity, per-unit discount, free product, or % "
    "discount) that automatically applies to Secondary Customer orders when the customer, sales "
    "channel, location and products all match."
)

doc.add_heading("The building blocks", level=3)
add_table(
    ["Object", "What it is", "Created via"],
    [
        ["Scheme Product Group", "A named, reusable list of products (SKUs) for a sales channel. Group-based schemes point at one of these.", "Scheme Product Group builder"],
        ["Scheme Product Group Item", "One product (SKU) inside a group.", "Added inside the group builder"],
        ["Scheme", "The offer header — name, type, sales channel, validity dates, and what it links to.", "Scheme creation wizard"],
        ["Scheme Slab", "A benefit tier under a scheme (e.g. 'Buy 10 → get 2 free').", "Scheme wizard (Slab Definition)"],
        ["Scheme Applicability", "Targeting rows — which Regions / Areas / Territories / Outlet Categories the scheme is valid for. Created automatically on Save.", "Scheme wizard (Applicability step)"],
    ],
    widths=[1.6, 3.4, 1.6],
)

doc.add_heading("The five Scheme Types", level=3)
add_table(
    ["Scheme Type", "What the customer gets", "Links to", "Slab based on"],
    [
        ["Free Quantity", "Free units of the same product group (price diluted across the free units).", "A Product Group", "Quantity"],
        ["QPS (Quantity Purchase Scheme)", "A flat ₹ per unit discount.", "A Product Group", "Quantity"],
        ["FOC Giveaway", "Free units of a different (free-of-cost) product.", "A Product Group", "Quantity"],
        ["Category Value", "A % discount on the value of a product sub-group/category.", "A Product Category", "Value (₹)"],
        ["Order Value", "A % discount on the whole order value.", "Nothing (whole order)", "Value (₹)"],
    ],
    widths=[1.4, 3.0, 1.4, 1.0],
)
para(
    "Rule of thumb: Free Quantity, QPS and FOC Giveaway need a Scheme Product Group. "
    "Category Value needs a Product Category. Order Value needs no product linkage.",
    italic=True, color=GREEN,
)

# ----------------------------------------------------------------------------- Part 1
doc.add_heading("Part 1 — Scheme Product Creation", level=1)
para(
    "A Scheme Product Group is the reusable bucket of SKUs that a group-based scheme (Free "
    "Quantity, QPS, FOC Giveaway) points at. Always create the group first, then reference it when "
    "you build the scheme in Part 2."
)

doc.add_heading("Process steps", level=2)
numbered([
    "Open the Scheme Product Group screen and click New (or open an existing group to edit it).",
    "Fill the header — Name, Group Purpose, Sales Channel, optional Description, and the Active "
    "toggle (see the field table below).",
    "Select the Sales Channel. The SKU table stays hidden ('Select a Sales Channel') until a "
    "channel is chosen, then loads the products for that channel.",
    "Narrow the list with the filter row (Category, Product Group, Sub Group, Variant, Grammage, "
    "MRP) or the Search SKU box (product name / SKU code).",
    "Tick the checkbox on each SKU you want in the group — the header shows a live 'n selected' "
    "count. Watch the conflict (⚠) flag.",
    "Click Save. The group and its selected products (Scheme Product Group Items) are stored "
    "together and the group becomes available in the Scheme wizard.",
])

doc.add_heading("Header fields", level=2)
add_table(
    ["Field", "Required", "Notes"],
    [
        ["Name", "Yes", "A clear, recognizable name, e.g. 'Pudding Cake 500 G'."],
        ["Group Purpose", "Yes", "Price Division — for Free Quantity / QPS groups. FOC Qualifier — for FOC Giveaway 'buy' groups. This is how the scheme wizard later filters which groups you can pick."],
        ["Sales Channel", "Yes", "The channel this group belongs to. The SKU list loads only after a channel is chosen."],
        ["Description", "No", "Optional notes."],
        ["Active", "—", "Toggle Yes to make the group usable; No keeps it as a draft."],
    ],
    widths=[1.3, 1.0, 4.3],
)

doc.add_heading("Filter & selection helpers", level=2)
bullets([
    "Clear Filters — reset all filters.",
    "Select All Visible — tick every row currently shown by the filters.",
    "Clear Selection — untick everything.",
    "All / Selected (n) toggle — switch between the full list and just your picked items.",
])
para("SKU columns: SKU Code, Product Name, Category, Group, Sub Group, Variant, Grammage, MRP, and a "
     "warning flag.")
para("Conflict warning: If a SKU already belongs to another group with the same purpose and channel, "
     "a warning icon appears (hover to see which group). Avoid putting the same SKU in two competing "
     "groups, or schemes can overlap unexpectedly.", italic=True, color=GREY)
para("Tip: Keep groups focused (e.g. one pack size / brand family). You'll reuse them across many "
     "schemes over time.", italic=True, color=GREEN)

# ----------------------------------------------------------------------------- Part 2
doc.add_heading("Part 2 — Scheme Creation and Applicability", level=1)
para("Schemes are built in a 2-step wizard: ① Definition and ② Applicability. Open the Scheme tab "
     "and click New (existing schemes open the same wizard in edit mode). Use Next / Back, and Save "
     "on the last step. Cancel discards.")

doc.add_heading("Process steps", level=2)
numbered([
    "Open the Scheme tab → New to launch the wizard.",
    "Step 1 — Master Details: enter the Scheme Name, Sales Channel, Scheme Type, Start/End dates, "
    "and the Active toggle.",
    "Step 1 — Linkage: depending on the Scheme Type, pick a Scheme Product Group (Free Quantity / "
    "QPS / FOC Giveaway), a Product Category (Category Value), or nothing (Order Value).",
    "Step 1 — Slab Definition: click Add Slab and fill the tier columns (they adapt to the Scheme "
    "Type). Add more slabs for multiple tiers; the trash icon removes a row. Click Next.",
    "Step 2 — Applicability: for each level (Outlet Categories, Regions, Areas, Territories) either "
    "leave 'Apply to all' ON or turn it OFF and move specific values into Selected.",
    "Click Save. The Scheme, all its Slabs, and the Applicability rows are saved together in one "
    "action. Re-open the scheme anytime to edit it.",
])

doc.add_heading("Step 1 — Definition", level=2)
para("Step 1 has three cards: Master Details, Product Group / Category, and Slab Definition.")

doc.add_heading("Master Details", level=3)
add_table(
    ["Field", "Required", "Notes"],
    [
        ["Scheme Name", "Yes", "Shown to order-takers, so make it descriptive (e.g. 'QPS Scheme – Pudding Cakes')."],
        ["Sales Channel", "Yes", "Drives which product groups/FOC products you can pick, and which customers see the scheme."],
        ["Scheme Type", "Yes", "Free Quantity / QPS / FOC Giveaway / Category Value / Order Value. This choice changes the linkage card and the slab columns."],
        ["Scheme Start Date", "Yes", "First day the scheme is live. Cannot be set in the past for a new scheme."],
        ["Scheme End Date", "Yes", "Last day the scheme is live (on/after the start date)."],
        ["Active", "—", "Toggle on to make it live (in addition to being within the date range)."],
        ["Description", "No", "Optional internal notes."],
    ],
    widths=[1.4, 1.0, 4.2],
)

doc.add_heading("Product Group / Category (linkage)", level=3)
para("What this card shows depends on the Scheme Type:")
bullets([
    "**Free Quantity / QPS / FOC Giveaway → a Scheme Product Group search. It lists active groups for your Sales Channel with the matching Group Purpose (Price Division for Free Qty/QPS, FOC Qualifier for FOC). Search by name and click to select; use Change to pick a different one.",
    "**Category Value → a Scheme Product Category dropdown. Pick the sub-group/category whose value drives the discount.",
    "**Order Value → no linkage needed ('Order Value schemes apply to the order's overall value').",
])
para("If you haven't chosen a Scheme Type and Sales Channel yet, this card shows 'Pick Scheme Type "
     "and Sales Channel first.'", italic=True, color=GREY)

doc.add_heading("Slab Definition (the tiers)", level=3)
para("Add one or more slabs (tiers). The columns adapt to the Scheme Type. Use Add Slab to add a "
     "row and the trash icon to remove one.")
add_table(
    ["Scheme Type", "Slab columns", "Meaning"],
    [
        ["Free Quantity", "Min Qty, Max Qty, Free Qty", "Buy within the qty band → get Free Qty free units."],
        ["QPS", "Min Qty, Max Qty, Benefit ₹/EA", "Buy within the qty band → ₹X off each unit."],
        ["FOC Giveaway", "Min Qty, Max Qty, FOC Product, Free Qty", "Buy within the qty band → get Free Qty units of the chosen FOC product, free."],
        ["Order Value", "Min Value (₹), Max Value (₹), Discount %", "Order value in the band → X% off the order."],
        ["Category Value", "Min Value (₹), Max Value (₹), Discount %", "Category value in the band → X% off that category."],
    ],
    widths=[1.3, 2.4, 2.9],
)
para("How tiers are chosen at order time: the system picks the highest tier the order qualifies for "
     "(by total quantity for qty-based schemes, or total value for value-based schemes).", bold=True)
para("Multi-tier example (Free Quantity):")
add_table(
    ["Min Qty", "Max Qty", "Free Qty"],
    [["10", "19", "1"], ["20", "(blank = no upper limit)", "3"]],
    widths=[1.2, 2.2, 1.2],
)
para("Ordering 25 EA → falls in the second tier → the engine grants free units based on that tier.")
para("FOC Product search: in an FOC slab, click the FOC Product cell, search by name or SKU code "
     "(scoped to the scheme's Sales Channel), and pick the giveaway product. Use the ✕ to change it.")

doc.add_heading("Step 2 — Applicability (targeting)", level=2)
para("Decide which customers the scheme reaches. Four levels, each with an 'Apply to all' toggle: "
     "Outlet Categories, Regions, Areas, Territories.")
bullets([
    "Leave Apply to all ON → the scheme matches every value at that level.",
    "Turn it OFF → a dual-list box appears; move the specific Available values into Selected to "
    "restrict the scheme to just those.",
])
para("Example: Apply to all Regions = ON, but Areas = OFF with only 'North Zone' selected → the "
     "scheme runs in every region but only for the North Zone area.", italic=True, color=GREY)

doc.add_heading("Save", level=3)
para("Click Save. The system saves the Scheme, all its Slabs, and the Applicability rows together "
     "in one action. Re-open the scheme anytime to edit it.")
para("Activation checklist — a scheme only goes live when ALL are true: Active = on, today is "
     "within Start/End dates, the Sales Channel matches, the Applicability matches the customer, "
     "and the linked group/category has products.", bold=True, color=GREEN)

doc.add_heading("Scheme Types — with an example each", level=2)
para("There are five scheme types. Each example below shows the slab you would enter and what the "
     "customer gets at order time.")

doc.add_heading("1) Free Quantity", level=3)
para("Free units of the same product group; the free units' value is spread (price diluted) across "
     "the line so the line total reflects the giveaway.")
para("Example — Group 'Pudding Cake 500 G', slab Min Qty 10, Max Qty 19, Free Qty 2. The customer "
     "orders 12 EA → qualifies for the 10–19 tier → gets 2 EA free. The unit price is diluted across "
     "the 12 units so the 2 free units cost nothing overall.")

doc.add_heading("2) QPS (Quantity Purchase Scheme)", level=3)
para("A flat ₹ per-unit discount for buying within a quantity band.")
para("Example — Group 'Pudding Cake 500 G', slab Min Qty 42, Max Qty 71, Benefit ₹1 / EA. The "
     "customer orders 50 EA at ₹20 → each unit is discounted ₹1 → pays ₹19 / EA (₹50 total saving).")

doc.add_heading("3) FOC Giveaway", level=3)
para("Free units of a different (free-of-cost) product when the 'buy' group reaches the band.")
para("Example — Buy group 'Whole Wheat Chakki Atta', slab Min Qty 20, FOC Product 'Roasted "
     "Vermicelli 200 g', Free Qty 2. The customer buys 20 EA of Atta → gets 2 EA of Roasted "
     "Vermicelli free. If Vermicelli is also ordered, the 2 free units merge into that line "
     "(shown as '+2 free'); otherwise a separate FREE line is added at ₹0.")

doc.add_heading("4) Category Value", level=3)
para("A % discount on the value of a product sub-group / category, shown as a header discount.")
para("Example — Category 'Plum', slab Min Value ₹10,000, Max Value ₹19,000, Discount 1%. The "
     "customer's Plum-category lines total ₹12,000 → qualifies → 1% (₹120) is taken off as a "
     "'Category Value Discount' on the order.")

doc.add_heading("5) Order Value", level=3)
para("A % discount on the whole order value (no product linkage), shown as a header discount and "
     "reflected in Net Payable.")
para("Example — slab Min Value ₹10,000, Max Value ₹19,000, Discount 1%. The order totals ₹15,000 → "
     "1% off the order. Order Value is applied after any Category Value discount (i.e. on the value "
     "remaining after the category discount).")

# ----------------------------------------------------------------------------- Part 3
doc.add_heading("Part 3 — Creation of Order", level=1)
para("When a salesperson creates a Secondary Customer order, matching schemes show up automatically "
     "and prices recalculate as quantities are entered. Schemes apply only for Secondary Customers — "
     "Primary Customers see base prices and a 'No schemes are applicable' message.")

doc.add_heading("Process steps", level=2)
numbered([
    "From the Beat Planner visit / Execute screen (or the Order object's New action), open the order "
    "screen for a Secondary Customer. The customer (account) is carried into the screen; if none is "
    "passed, use the Search Customer picker.",
    "Products screen ('All items'): search / filter products and enter the EA quantity for each. A "
    "⭐ / ribbon marks scheme products; an applied line shows a discounted price band next to the "
    "original price, with a breakup you can expand.",
    "Schemes ('Running Schemes') tab: review the applicable schemes (read-only) — each card shows the "
    "name, type, offer tiers as sentences, validity, and the covered group/category. Expand / collapse "
    "with the chevron.",
    "Selected tab: quick-review every product that has a non-zero quantity before finishing.",
    "Summary screen: enter Delivery To and Expected Delivery Date (PO Number / PO Date appear for "
    "Primary Customers only). Review per-line discounted prices, Sub Total, Tax, Grand Total, any "
    "header discounts (Category Value / Order Value) and the resulting Net Payable. FOC / free lines "
    "are tagged FREE (or '+n free' when merged).",
    "Click Save Orders. The system validates the required header fields and the minimum order value, "
    "captures the device location (GPS), then saves the order.",
    "After save: a 'Orders saved successfully' toast appears and you return to the visit / Execute "
    "screen (or the Order list view if you came from the Order New action).",
])
para("Every applied benefit is also recorded against the order for auditing / claims (Order Scheme "
     "Applied rows: which scheme, which tier, the benefit amount, free quantity, etc.).", italic=True,
     color=GREY)

doc.add_heading("Which schemes appear (coverage rules)", level=2)
para("A scheme appears for a customer when all of the following hold:")
numbered([
    "The account is a Secondary Customer (primary customers don't see these schemes).",
    "The scheme's Sales Channel matches the customer's channel.",
    "Today is within the scheme's Start–End dates and the scheme is Active.",
    "The customer's Region / Area / Territory match the scheme's applicability (or 'Apply to all').",
    "The customer's Outlet Category matches (or 'Apply to all Outlet Categories').",
    "The scheme has product membership — a product group, a category, or it's an Order Value scheme.",
])

doc.add_heading("The 'Running Schemes' list", level=2)
para("The order screen shows a Running Schemes panel of cards — one per applicable scheme. Each card shows:")
bullets([
    "Scheme name and Scheme Type.",
    "Offer tiers as plain sentences, e.g. Free Quantity → '🛒 Buy 10 EA → 🎁 Get 2 EA Free'; "
    "QPS → '🛒 Buy 42 to 71 EA → 💰 ₹1 off per EA'; FOC → '🛒 Buy 20 EA → 🎁 Get 3 EA of (product) "
    "Free'; Category/Order Value → '🛒 Order ₹10000 to ₹19000 → 💰 1% Off'.",
    "Multiple tiers are labelled Tier 1, Tier 2, …; a single tier is labelled Offer.",
    "Validity — e.g. 31-05-2026 to 30-06-2026.",
    "Group + member products (name + SKU) for group schemes, or the Category for category schemes.",
])
para("(In the invoice wizard, each card can be expanded/collapsed with the chevron.)", italic=True, color=GREY)

doc.add_heading("Indicators on each product row", level=2)
bullets([
    "A ribbon / ⭐ icon marks products that belong to a running scheme.",
    "When a scheme applies to a line, a price band shows the discounted unit price next to the "
    "original price, plus a tag with the scheme name. A breakup control reveals the step-by-step "
    "price changes (Base → Free Quantity → QPS → …).",
])

doc.add_heading("What each scheme type does to the price", level=2)
para("As quantities are entered, the engine recalculates in a fixed order. Multiple schemes can "
     "apply to the same products.")
bullets([
    "**Free Quantity — totals the group's quantity, works out the free units for the qualifying "
    "tier, and dilutes the unit price so the free units are reflected.",
    "**QPS — subtracts the ₹/EA benefit from each unit's price for the qualifying tier.",
    "**FOC Giveaway — adds the free product. If already on the order, the free units are merged "
    "into that line (shown as '+n free'); otherwise a new FREE line is added at zero net cost.",
    "**Category Value — % discount on the category's value, shown as a header discount.",
    "**Order Value — % discount on the whole order, shown as a header discount.",
])
para("For value-based schemes, the system always uses the highest tier the order value qualifies for.")

doc.add_heading("The order summary", level=2)
bullets([
    "FOC / free lines are tagged FREE (or a line shows '+n free' for merged units).",
    "Per-line prices show the discounted unit price.",
    "Header discounts appear when present: Category Value Discount – ₹…, Order Value Discount – ₹…, "
    "and Net Payable = Grand Total (subtotal + tax) − header discounts.",
])
para("Every applied benefit is also recorded against the order for auditing/claims (which scheme, "
     "which tier, the benefit amount, free quantity, etc.).")

# ----------------------------------------------------------------------------- Appendix
doc.add_heading("Appendix — Troubleshooting: 'Why isn't my scheme showing?'", level=1)
numbered([
    "Active & dates — Is the scheme Active, and is today between its Start and End dates?",
    "Customer type — Is the order for a Secondary Customer? Primary customers don't see these.",
    "Sales channel — Does the scheme's Sales Channel match the customer's channel?",
    "Applicability — Do the customer's Region / Area / Territory / Outlet Category match? If you "
    "restricted any level, confirm the customer's value is in the Selected list (or switch to 'Apply to all').",
    "Product linkage — Group schemes: does the Scheme Product Group contain the ordered products, "
    "and is it Active with the right Group Purpose and Sales Channel? Category Value: is the Scheme "
    "Product Category correct and are products from it on the order?",
    "Quantity / value threshold — Has the order reached a slab's Min Qty or Min Value? Below the "
    "lowest tier, nothing applies. (In the invoice wizard, reducing quantity below the threshold "
    "removes the benefit and flags it as a scheme issue.)",
    "Group conflicts — Was the SKU placed in two groups of the same purpose/channel (the ⚠ flag in "
    "the group builder)? Clean up duplicates.",
])

doc.add_heading("Quick reference — which fields matter per Scheme Type", level=2)
add_table(
    ["Scheme Type", "Linkage", "Slab inputs", "Result"],
    [
        ["Free Quantity", "Product Group", "Min/Max Qty, Free Qty", "Free units of same group"],
        ["QPS", "Product Group", "Min/Max Qty, Benefit ₹/EA", "₹ off per unit"],
        ["FOC Giveaway", "Product Group", "Min/Max Qty, FOC Product, Free Qty", "Free units of another product"],
        ["Category Value", "Product Category", "Min/Max Value, Discount %", "% off the category value"],
        ["Order Value", "(none)", "Min/Max Value, Discount %", "% off the whole order"],
    ],
    widths=[1.3, 1.4, 2.2, 1.7],
)

import os
out = os.path.join(os.path.dirname(__file__), "Schemes-User-Manual.docx")
doc.save(out)
print("Saved", out)
